import os
from fpdf import FPDF
from docx import Document
from docx.shared import Pt
from Management.models import Attendance, Course, Student


class ReportGenerator:
    """
    A class to generate reports in DOCX and PDF formats based on attendance data.

    Attributes:
        report_data (str): The name of the report data.
        output_dir (str): The name of the output directory.
    """

    def __init__(self, report_data_name='report', output_dir_name='output'):
        """
        Initializes the ReportGenerator with the given report data name and output directory name.

        Args:
            report_data_name (str): The name of the report data.
            output_dir_name (str): The name of the output directory.
        """
        self.report_data = report_data_name
        self.output_dir = output_dir_name

    def create_word_files(self, room_data, term, exam_type):
        """
        Creates Word documents for each course based on the provided room data, term, and exam type.

        Args:
            room_data (list): A list of dictionaries containing room and student data.
            term (str): The term for which the report is generated.
            exam_type (str): The type of exam.

        Returns:
            list: A list of file paths to the generated Word documents.
        """
        course_docs = {}
        output_dir = self.output_dir
        file_paths = []

        # Ensure output directory exists
        os.makedirs(output_dir, exist_ok=True)

        def add_paragraph(doc, text, bold=True, font_size=12):
            """ 
            Helper function to add a formatted paragraph to the document. 

            Args:
                doc (Document): The Word document object.
                text (str): The text to add to the paragraph.
                bold (bool): Whether the text should be bold.
                font_size (int): The font size of the text.
            """
            paragraph = doc.add_paragraph()
            run = paragraph.add_run(text)
            run.bold = bold
            run.font.size = Pt(font_size)

        for room in room_data:
            room_no = room['ROOM_NO']

            for course in room['STUDENT_LIST']:
                course_code = course['course_code']
                course_data = Course.objects.filter(
                    COURSE_CODE=course_code, TERM=term).first()

                # Fetch or create the document for the course
                doc = course_docs.get(course_code, Document())
                course_docs[course_code] = doc

                # Insert a page break if the document already has content
                if len(doc.paragraphs) > 0:
                    doc.add_page_break()

                # Add exam information, course code, and room number
                add_paragraph(doc, f"\t\tExam:\t{term} / {exam_type}")
                add_paragraph(doc, f'\t\tCourse:\t{course_code}')
                add_paragraph(doc, f"\t\tCourse Name:\t{
                              course_data.COURSE_NAME}")
                add_paragraph(doc, f"\t\tRoom:\t{room_no}")

                # Create or add to the table
                table = doc.add_table(rows=1, cols=3)
                table.style = 'Table Grid'

                # Set the table headers
                hdr_cells = table.rows[0].cells
                hdr_cells[0].text = 'No.'
                hdr_cells[1].text = 'Name'
                hdr_cells[2].text = 'Batch'

                # Populate table with student data
                for idx, student_id in enumerate(course['students'], start=1):
                    student = Student.objects.get(STUDENT_ID=student_id)
                    row_cells = table.add_row().cells
                    row_cells[0].text = str(idx)
                    row_cells[1].text = student.STUDENT_NAME
                    row_cells[2].text = student.STUDENT_BATCH

        # Save each course's document after all rooms are added
        for course_code, doc in course_docs.items():
            file_path = os.path.join(output_dir, f"{course_data.COURSE_NAME}({course_code}) {
                                     exam_type} - Sitting Arrangement.docx")
            doc.save(file_path)
            file_paths.append(file_path)

        return file_paths

    def generate_report(self, criteria, value, term, exam_type, output_dir='output', file_type='docx'):
        """
        Generates a report based on the given criteria and value, and saves it in the specified file type.

        Args:
            criteria (str): The criteria to filter attendance data ('ROOM_NO', 'COURSE_CODE', or 'STUDENT_BATCH').
            value (str): The value for the specified criteria.
            output_dir (str): The directory to save the generated report.
            file_type (str): The type of file to generate ('docx' or 'pdf').

        Returns:
            str: The file path to the generated report.
        """

        # Query the Attendance model based on the criteria
        if criteria == 'ROOM_NO':
            attendances = Attendance.objects.filter(
                ROOM_NO__ROOM_NO__iexact=value)
            attendances = attendances.filter(
                EXAM_ID__COURSE_CODE__TERM__iexact=term)
            attendances = attendances.filter(
                EXAM_ID__EXAM_TYPE__iexact=exam_type)
        elif criteria == 'COURSE_CODE':
            attendances = Attendance.objects.filter(
                EXAM_ID__COURSE_CODE__COURSE_CODE__iexact=value)
            attendances = attendances.filter(
                EXAM_ID__COURSE_CODE__TERM__iexact=term)
            attendances = attendances.filter(
                EXAM_ID__EXAM_TYPE__iexact=exam_type)
        elif criteria == 'STUDENT_BATCH':
            attendances = Attendance.objects.filter(
                STUDENT_ID__STUDENT_BATCH__iexact=value)
            attendances = attendances.filter(
                EXAM_ID__COURSE_CODE__TERM__iexact=term)
            attendances = attendances.filter(
                EXAM_ID__EXAM_TYPE__iexact=exam_type)
        else:
            raise ValueError(
                "Invalid criteria. Choose from 'ROOM_NO', 'COURSE_CODE', or 'STUDENT_BATCH'.")

        # Define header columns for the report
        headers = ['No.', 'Student Name', 'Course Code',
                   'Room No.', 'Attendance Status']

        def populate_data(data):
            """ 
            Helper function to map attendance data into rows. 

            Args:
                data (QuerySet): The attendance data to map.

            Returns:
                list: A list of lists containing the mapped attendance data.
            """
            return [
                [str(idx), attendance.STUDENT_ID.STUDENT_NAME,
                 attendance.EXAM_ID.COURSE_CODE.COURSE_CODE,
                 attendance.ROOM_NO.ROOM_NO,
                 'Present' if attendance.ATTENDANCE_STATUS else 'Absent']
                for idx, attendance in enumerate(data, start=1)
            ]

        # Ensure the output directory exists
        os.makedirs(output_dir, exist_ok=True)

        if file_type == 'docx':
            # Same logic for docx generation

            doc = Document()

            # Add title to the document
            title = doc.add_heading(level=1).add_run(
                f'Attendance Report by {criteria}: {value}')
            title.bold = True
            title.font.size = Pt(16)

            # Create a table and add headers
            table = doc.add_table(rows=1, cols=len(headers))
            table.style = 'Table Grid'
            hdr_cells = table.rows[0].cells
            for i, header in enumerate(headers):
                hdr_cells[i].text = header

            # Populate the table with attendance data
            for row_data in populate_data(attendances):
                row_cells = table.add_row().cells
                for i, cell_text in enumerate(row_data):
                    row_cells[i].text = cell_text

            # Save the document
            file_path = os.path.join(output_dir, f"Attendance_Report_{
                                     criteria}_{value}.docx")
            doc.save(file_path)

        elif file_type == 'pdf':
            # Use fpdf for PDF generation
            file_path = os.path.join(output_dir, f"Attendance_Report_{
                                     criteria}_{value}.pdf")

            pdf = FPDF()
            pdf.set_auto_page_break(auto=True, margin=15)
            pdf.add_page()

            # Set fonts and title
            pdf.set_font("Arial", 'B', 16)
            pdf.cell(200, 10, f'Attendance Report for: {
                     value}', ln=True, align='C')

            # Add a header row with adjusted column widths
            pdf.set_font("Arial", 'B', 8)

            # Adjust the column widths: No. is slimmer, Name has more space
            # Custom widths for No., Name, Course Code, Room No., Attendance Status
            col_widths = [20, 90, 20, 20, 30]

            headers = ['No.', 'Student Name', 'Course Code',
                       'Room No.', 'Attendance Status']
            for i, header in enumerate(headers):
                pdf.cell(col_widths[i], 8, header, 1, 0, 'C')
            pdf.ln()

            # Populate the table with attendance data and adjusted column widths
            pdf.set_font("Arial", '', 12)
            for row_data in populate_data(attendances):
                for i, cell_text in enumerate(row_data):
                    pdf.cell(col_widths[i], 10, cell_text, 1)
                pdf.ln()

            # Save the PDF
            pdf.output(file_path)

        else:
            raise ValueError("Invalid file type. Choose 'docx' or 'pdf'.")

        return file_path
