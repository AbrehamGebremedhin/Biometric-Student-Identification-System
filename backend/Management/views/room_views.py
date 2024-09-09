import os
import shutil
import zipfile
from io import BytesIO
import pandas as pd
from Management.models import Room, Course, Exam, Student
from Management.serializers import RoomSerializer
from django.http import Http404
from rest_framework.views import APIView
from rest_framework.response import Response
from rest_framework import status
from rest_framework.permissions import IsAuthenticated
from docx import Document
from docx.shared import Pt
from django.http import FileResponse


class RoomList(APIView):
    def get(self, request):
        room_no = request.query_params.get('room_no')
        exam_time = request.query_params.get('exam_time')

        if not room_no and not exam_time:
            rooms = Room.objects.all()
            serializer = RoomSerializer(rooms, many=True)
            return Response(serializer.data)

        if room_no:
            rooms = rooms.filter(ROOM_NO__icontains=room_no)

        if exam_time:
            rooms = rooms.filter(EXAM_TIME__icontains=exam_time)

        if not rooms.exists():
            return Response({"error": "Room not found"}, status=status.HTTP_404_NOT_FOUND)

        serializer = RoomSerializer(rooms, many=True)
        return Response(serializer.data)

    def post(self, request):
        file = request.FILES.get('file')
        if file.name.endswith('.csv'):
            data = pd.read_csv(file)
        elif file.name.endswith('.xlsx') or file.name.endswith('.xls'):
            data = pd.read_excel(file)
        else:
            return Response({"error": "Unsupported file format"}, status=status.HTTP_400_BAD_REQUEST)

        # Required columns
        required_columns = {'room_no', 'student_id', 'course_code', 'term'}

        # Check if all required columns are present
        if not required_columns.issubset(data.columns):
            missing_columns = required_columns - set(data.columns)
            return Response({"error": f"Missing columns: {', '.join(missing_columns)}"}, status=status.HTTP_400_BAD_REQUEST)

        term = data['term'].iloc[0] if 'term' in data.columns else 'Unknown Term'
        unique_room_nos = data['room_no'].unique()
        room_data = []

        for room_no in unique_room_nos:
            students = []
            course_codes = data[data['room_no'] ==
                                room_no]['course_code'].unique()

            for course_code in course_codes:
                course = Course.objects.filter(
                    COURSE_CODE=course_code, TERM=term).first()
                if not course:
                    return Response({"error": f"Course {course_code} in {term} does not exist"}, status=status.HTTP_400_BAD_REQUEST)
                exam = Exam.objects.filter(COURSE_CODE=course).first()

                if not exam:
                    return Response({"error": f"Exam {course_code} in {term} does not exist"}, status=status.HTTP_400_BAD_REQUEST)

                exam_type = exam.EXAM_TYPE

                student_ids = data[(data['room_no'] == room_no) & (
                    data['course_code'] == course_code)]['student_id'].unique().tolist()

                for student_id in student_ids:
                    try:
                        Student.objects.get(STUDENT_ID=student_id)
                    except Student.DoesNotExist:
                        return Response({"error": f"Student {student_id} does not exist"}, status=status.HTTP_400_BAD_REQUEST)

                students.append({
                    "course_code": course_code,
                    "students": list(student_ids)
                })

            room_data.append({
                "ROOM_NO": room_no,
                "EXAM_TIME": request.data.get("EXAM_TIME"),
                "STUDENT_LIST": students
            })

        serializer = RoomSerializer(data=room_data, many=True)

        if serializer.is_valid():
            serializer.save()
            # Create the word files and get the output directory
            output_dir = self.create_word_files(room_data, term, exam_type)

            # Zip the output directory
            zip_buffer = BytesIO()
            with zipfile.ZipFile(zip_buffer, 'w') as zip_file:
                for root, dirs, files in os.walk(output_dir):
                    for file in files:
                        file_path = os.path.join(root, file)
                        zip_file.write(file_path, os.path.relpath(
                            file_path, output_dir))

            # Set buffer position to the start
            zip_buffer.seek(0)

            # Clean the output directory
            shutil.rmtree(output_dir)

            return FileResponse(zip_buffer, as_attachment=True, filename="exam_documents.zip")
        return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)

    def create_word_files(self, room_data, term, exam_type):
        course_docs = {}
        file_paths = []

        # Define the output directory
        output_dir = 'output'
        os.makedirs(output_dir, exist_ok=True)

        for room in room_data:
            room_no = room['ROOM_NO']
            for course in room['STUDENT_LIST']:
                course_code = course['course_code']
                course_data = Course.objects.filter(
                    COURSE_CODE=course_code, TERM=term).first()

                if course_code not in course_docs:
                    doc = Document()
                    course_docs[course_code] = doc
                else:
                    doc = course_docs[course_code]

                # Insert a page break before adding a new room section if it's not the first room
                if len(doc.paragraphs) > 0:
                    doc.add_page_break()

                # Reformat the data above each table
                exam_paragraph = doc.add_paragraph()
                exam_run = exam_paragraph.add_run(
                    f"Exam: {term} / {exam_type}")
                exam_run.bold = True
                exam_run.font.size = Pt(14)

                course_code_paragraph = doc.add_paragraph()
                course_code_run = course_code_paragraph.add_run(
                    f'Course: {course_code}')
                course_code_run.bold = True
                course_code_run.font.size = Pt(14)

                course_name_paragraph = doc.add_paragraph()
                course_name_run = course_name_paragraph.add_run(
                    f"Course Name: {course_data.COURSE_NAME}")
                course_name_run.bold = True
                course_name_run.font.size = Pt(14)

                room_paragraph = doc.add_paragraph()
                room_run = room_paragraph.add_run(f"Room: {room_no}")
                room_run.bold = True
                room_run.font.size = Pt(14)

                # Create a new table for each room in the same document
                table = doc.add_table(rows=1, cols=3)
                table.style = 'Table Grid'
                hdr_cells = table.rows[0].cells
                hdr_cells[0].text = 'No.'
                hdr_cells[1].text = 'Name'
                hdr_cells[2].text = 'Batch'

                for idx, student_id in enumerate(course['students'], start=1):
                    student = Student.objects.get(STUDENT_ID=student_id)
                    row_cells = table.add_row().cells
                    row_cells[0].text = str(idx)
                    row_cells[1].text = student.STUDENT_NAME
                    row_cells[2].text = student.STUDENT_BATCH

            # Save each course's document after adding all rooms
            for course_code, doc in course_docs.items():
                file_path = os.path.join(output_dir, f"{course_code}.docx")
                doc.save(file_path)

        return output_dir


class RoomDetail(APIView):
    """
    Retrieve, update or delete an exam instance.
    """
    permission_classes = [IsAuthenticated]

    def get_object(self, no):
        try:
            return Room.objects.get(ROOM_NO=no)
        except Room.DoesNotExist:
            raise Http404

    def get(self, request, no):
        room = self.get_object(no)
        serializer = RoomSerializer(room)
        return Response(serializer.data)

    def delete(self, request, no):
        room = self.get_object(no)
        room.delete()
        return Response(data={True}, status=status.HTTP_204_NO_CONTENT)
